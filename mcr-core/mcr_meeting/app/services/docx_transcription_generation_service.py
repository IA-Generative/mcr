import os
from abc import ABC, abstractmethod
from io import BytesIO
from typing import Any, List, Optional

from docx import Document as CreateDocument
from docx.document import Document as Docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from mcr_meeting.app.models import Transcription


class TemplatedDocxGenerator(ABC):
    doc: Docx

    def __init__(self, filename: str):
        decision_template_path = os.path.join(
            os.getcwd(), "mcr_meeting", "app", "cr-templates", filename
        )
        if not os.path.exists(decision_template_path):
            raise FileNotFoundError(
                f"Template file not found at {decision_template_path}"
            )

        self.doc = CreateDocument(decision_template_path)

    @abstractmethod
    def fill_templated_doc(self, *args: Any, **kwargs: Any) -> None:  # type: ignore[explicit-any]
        pass

    def save_and_return_docx(self) -> BytesIO:
        docx_io = BytesIO()
        self.doc.save(docx_io)
        docx_io.seek(0)  # Reset the buffer position to the beginning
        return docx_io

    def set_title(self, title: str) -> None:
        TITLE_PLACEHOLDER = "{{meeting_name}}"
        for para in self.doc.paragraphs:
            if TITLE_PLACEHOLDER in para.text:
                title_para = para
                break

        else:
            raise ValueError(f"Couldn't find {TITLE_PLACEHOLDER} in template")

        title_para.clear()  # type: ignore[no-untyped-call]
        run = title_para.add_run()
        run.text = title


class TranscriptionDocxGenerator(TemplatedDocxGenerator):
    def fill_templated_doc(
        self, meeting_name: str, transcriptions: List[Transcription]
    ) -> None:
        self.set_title(meeting_name)

        for t in transcriptions:
            para = self.doc.add_paragraph()
            self.add_speaker(paragraph=para, speaker=t.speaker)
            self.add_transcription_text(
                paragraph=para, transcription_text=t.transcription
            )
            self.format_transcription_paragraph(paragraph=para)

    def add_speaker(self, paragraph: Paragraph, speaker: str) -> None:
        run = paragraph.add_run(f"{speaker} : ")
        run.bold = True

    def add_transcription_text(
        self, paragraph: Paragraph, transcription_text: str
    ) -> None:
        paragraph.add_run(transcription_text)

    def format_transcription_paragraph(self, paragraph: Paragraph) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(
            12
        )  # 12 points space after this paragraph


def generate_transcription_docx(
    meeting_name: Optional[str], transcriptions: List[Transcription]
) -> BytesIO:
    """
    Generates a DOCX document containing the transcription of a meeting in a tabular format.

    Args:
        meeting_name (str): The name of the meeting to include in the document heading.
        transcriptions (list): A list of transcription objects, where each object contains `speaker` and `transcription`.

    Returns:
        BytesIO: A memory buffer containing the generated DOCX file.
    """
    doc_generator = TranscriptionDocxGenerator("FCR_transcription_template.docx")
    title = meeting_name if meeting_name is not None else "Transcription"
    doc_generator.fill_templated_doc(title, transcriptions)
    return doc_generator.save_and_return_docx()
