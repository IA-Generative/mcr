from pathlib import Path

from docx import Document


def load_docx_text(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def save_text_as_docx(text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    document.save(str(path))
